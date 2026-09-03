"""
Shared utilities for generating .docx reports (Technical Report and
Individual Contribution Reports) that reuse the exact MSU Cover Page
template (logo, table borders, merged cells) instead of rebuilding it
from scratch.

Formatting rules applied to body content:
- Times New Roman, 12pt, black font
- 1.5 line spacing
- Tables labelled at the top, figures labelled at the bottom
- No em dashes
"""

import copy

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

COVER_PAGE_PATH = (
    r"C:\Users\NgonidzasheMuzanenha\OneDrive\Masters Information Systems"
    r"\Semester 1.2\Cover Page.docx"
)

BLACK = RGBColor(0, 0, 0)
FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)


def _set_cell_value(cell, new_text, bold=None):
    """Replace the text of a cell while preserving the first run's formatting."""
    paragraph = cell.paragraphs[0]
    runs = paragraph.runs
    if not runs:
        run = paragraph.add_run(new_text)
    else:
        runs[0].text = new_text
        if bold is not None:
            runs[0].font.bold = bold
        # Clear any additional runs so old fragments do not linger
        for extra in runs[1:]:
            extra.text = ""


def _find_row_by_label(table, label):
    """Return the row whose first cell text matches the given label."""
    for row in table.rows:
        if row.cells[0].text.strip() == label:
            return row
    return None


def build_cover_page_document(student_name, reg_number, question_text, module_code="MIM736"):
    """
    Load the original MSU Cover Page template and update it in place,
    preserving the logo image, table borders, and merged cells exactly
    as designed. Returns the Document object with the cover page as the
    first page, ready to have body content appended after a page break.
    """
    doc = Document(COVER_PAGE_PATH)

    # --- Table 0: Module code header block ---
    # First paragraph in this cell reads "MIM737 - SOFTWARE ENGINEERING" split
    # across several runs. The whole heading is rebuilt on the first run
    # (preserving its bold Times New Roman formatting) and the remaining
    # runs are cleared. The second paragraph "ASSIGNMENT" is left untouched.
    header_table = doc.tables[0]
    header_cell = header_table.rows[0].cells[0]
    header_para = header_cell.paragraphs[0]
    runs = header_para.runs
    if runs:
        runs[0].text = f"{module_code} \u2013 SOFTWARE ENGINEERING"
        for extra in runs[1:]:
            extra.text = ""

    # --- Table 1: Student/module details ---
    details_table = doc.tables[1]

    row = _find_row_by_label(details_table, "Student Name:")
    if row:
        _set_cell_value(row.cells[1], student_name)

    row = _find_row_by_label(details_table, "Registration Number:")
    if row:
        _set_cell_value(row.cells[1], reg_number)

    row = _find_row_by_label(details_table, "Module Code:")
    if row:
        _set_cell_value(row.cells[1], module_code)

    row = _find_row_by_label(details_table, "Module Name:")
    if row:
        _set_cell_value(row.cells[1], "Software Engineering")

    row = _find_row_by_label(details_table, "Lecturer:")
    if row:
        _set_cell_value(row.cells[1], "Dr Zhou")

    row = _find_row_by_label(details_table, "QUESTION:")
    if row:
        # Merged cell: append the question text on a new paragraph
        cell = row.cells[0]
        para = cell.add_paragraph()
        run = para.add_run(question_text)
        run.font.bold = False

    # Page break to separate cover page from body content
    doc.add_page_break()

    return doc


def set_document_defaults(doc):
    """Set the default Normal style to Times New Roman 12pt, black, 1.5 spacing."""
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    style.font.color.rgb = BLACK
    style.paragraph_format.line_spacing = 1.5
    # Ensure east-asian font is consistent too
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)


def add_heading(doc, text, level=1):
    """Add a heading formatted in Times New Roman, black."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(14 if level == 1 else 12)
        run.font.color.rgb = BLACK
        run.font.bold = True
    return heading


def add_paragraph(doc, text):
    """Add a body paragraph with 1.5 spacing, Times New Roman 12pt, black."""
    para = doc.add_paragraph(text)
    para.paragraph_format.line_spacing = 1.5
    for run in para.runs:
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
        run.font.color.rgb = BLACK
    return para


def _apply_grid_borders(table):
    """
    Apply visible grid borders to a table via raw OXML. Used as a fallback
    when the "Table Grid" named style is not present in the document (the
    MSU cover page template does not define it), so tables in the body of
    the report still render with borders.
    """
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = borders.makeelement(qn(f"w:{edge}"), {})
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def add_table(doc, label, headers, rows):
    """Add a table with its label ABOVE the table, per formatting rules."""
    label_para = doc.add_paragraph(label)
    label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_para.paragraph_format.line_spacing = 1.5
    for run in label_para.runs:
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
        run.font.bold = True
        run.font.color.rgb = BLACK

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        _apply_grid_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.line_spacing = 1.5
            for run in paragraph.runs:
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE
                run.font.bold = True
                run.font.color.rgb = BLACK

    for i, row_values in enumerate(rows):
        for j, value in enumerate(row_values):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.5
                for run in paragraph.runs:
                    run.font.name = FONT_NAME
                    run.font.size = FONT_SIZE
                    run.font.color.rgb = BLACK

    doc.add_paragraph()
    return table


def add_figure_label(doc, label):
    """Add a figure label BELOW the figure, per formatting rules."""
    label_para = doc.add_paragraph(label)
    label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_para.paragraph_format.line_spacing = 1.5
    for run in label_para.runs:
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
        run.font.bold = True
        run.font.color.rgb = BLACK
    return label_para


def finalize_document_fonts(doc):
    """
    Force Times New Roman 12pt black on every run in the document
    (paragraphs and tables), except within the cover page tables which
    already carry their own correct formatting from the template.
    This only touches runs that do not already specify a font, to avoid
    disturbing the cover page's original design.
    """
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing or 1.5
        for run in paragraph.runs:
            if run.font.name is None:
                run.font.name = FONT_NAME
            if run.font.size is None:
                run.font.size = FONT_SIZE

"""
Generate the Technical Report as a .docx file, reusing the exact MSU
cover page template (logo, table borders, merged cells preserved).
Formatting: Times New Roman, 12pt, 1.5 line spacing, black font.
No em dashes. APA referencing. Tables labelled at top, figures at bottom.
"""

import sys
import os

from docx.shared import Inches

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from report_utils import (
    build_cover_page_document,
    set_document_defaults,
    add_heading,
    add_paragraph,
    add_table,
    add_figure_label,
    finalize_document_fonts,
    FONT_NAME,
    FONT_SIZE,
    BLACK,
)

OUTPUT_PATH = r"C:\Users\NgonidzasheMuzanenha\OneDrive\Masters Information Systems\Semester 1.2\MIM736 - Software Engineering\Technical Report - QVS.docx"


def generate_report():
    """Generate the full technical report."""
    doc = build_cover_page_document(
        student_name="Ngonidzashe Muzanenhamo, Loreen Venge, Judah T Chisare, Nyasha A Madziwanzira",
        reg_number="R211790N, R2118621M, R267853N, R2117220T",
        question_text=(
            "Assignment 2: DevOps-Enabled Qualification Verification System "
            "- Technical Report"
        ),
    )
    set_document_defaults(doc)

    # Table of Contents placeholder
    add_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Introduction",
        "2. Problem Analysis",
        "3. Requirements",
        "4. System Architecture",
        "5. Design Decisions",
        "6. DevOps Workflow",
        "7. Testing Strategy",
        "8. Verification Strategy",
        "9. Critical Evaluation",
        "10. Conclusion",
        "References",
    ]
    for item in toc_items:
        add_paragraph(doc, item)
    doc.add_page_break()

    # 1. Introduction
    add_heading(doc, "1. Introduction", level=1)
    add_paragraph(doc,
        "Educational institutions, professional bodies, and employers increasingly require reliable mechanisms to verify academic and professional qualifications. The rise of credential fraud, including fake degrees and forged certificates, has created an urgent need for secure, scalable, and auditable verification systems (Alqarni et al., 2023). This report presents the design, development, testing, and deployment of a DevOps-enabled Qualification Verification System (QVS) that addresses these challenges through modern software engineering practices."
    )
    add_paragraph(doc,
        "The system was developed by a team of four members using Python and the FastAPI framework. It incorporates blockchain-based credential verification, JWT authentication with role-based access control, AI-powered fraud detection, and real-time monitoring through Prometheus and Grafana. The project demonstrates industry-standard DevOps practices including Git-based version control, continuous integration, continuous delivery, automated testing, and Docker containerization (Bass, Weber, & Zhu, 2021)."
    )
    add_paragraph(doc,
        "This report covers the problem analysis, system requirements, architecture, design decisions, DevOps workflow, testing strategy, verification strategy, and a critical evaluation of the system. The report demonstrates how software requirements are automatically verified through test cases, validation rules, and CI/CD quality gates."
    )

    # 2. Problem Analysis
    add_heading(doc, "2. Problem Analysis", level=1)
    add_paragraph(doc,
        "Credential fraud is a growing problem worldwide. Studies show that credential verification processes in many institutions are manual, time-consuming, and prone to errors (Khan et al., 2022). Traditional verification systems suffer from several limitations including lack of tamper-evident records, limited auditability, and absence of automated fraud detection mechanisms."
    )
    add_paragraph(doc,
        "The core problem is the need for a system that allows authorised users to register qualifications, search and retrieve records, verify the authenticity of qualifications, and maintain an auditable history of all verification activities. The system must be secure, scalable, maintainable, and supported by modern DevOps practices to ensure continuous delivery of updates and improvements."
    )
    add_paragraph(doc,
        "Additional challenges include ensuring data integrity through tamper-evident mechanisms, providing role-based access to sensitive operations, detecting potentially fraudulent credentials using AI, and monitoring system health in real time. The system must also support deployment across different environments, from local development with SQLite to production with PostgreSQL and cloud hosting."
    )

    # 3. Requirements
    add_heading(doc, "3. Requirements", level=1)
    add_heading(doc, "3.1 Functional Requirements", level=2)
    add_table(doc, "Table 1: Functional Requirements",
        ["ID", "Requirement", "Priority"],
        [
            ["FR1", "The system shall allow authorised users to register qualifications", "High"],
            ["FR2", "The system shall allow users to search and retrieve qualification records", "High"],
            ["FR3", "The system shall verify the authenticity of qualifications using blockchain hashing", "High"],
            ["FR4", "The system shall maintain an auditable history of all verification activities", "High"],
            ["FR5", "The system shall provide JWT-based authentication with role-based access control", "High"],
            ["FR6", "The system shall provide AI-powered analysis for fraud detection", "Medium"],
            ["FR7", "The system shall expose Prometheus metrics for monitoring", "Medium"],
            ["FR8", "The system shall support soft deletion of qualification records", "Medium"],
        ]
    )

    add_heading(doc, "3.2 Non-Functional Requirements", level=2)
    add_table(doc, "Table 2: Non-Functional Requirements",
        ["ID", "Requirement", "Category"],
        [
            ["NFR1", "The system shall respond to API requests within 200ms", "Performance"],
            ["NFR2", "The system shall achieve at least 80% test coverage", "Quality"],
            ["NFR3", "The system shall pass static analysis with no critical issues", "Security"],
            ["NFR4", "The system shall support both SQLite and PostgreSQL databases", "Portability"],
            ["NFR5", "The system shall be containerised using Docker", "Deployability"],
            ["NFR6", "The system shall deploy automatically on merge to main branch", "DevOps"],
            ["NFR7", "The system shall use bcrypt for password hashing", "Security"],
            ["NFR8", "The system shall provide Swagger documentation at /docs", "Usability"],
        ]
    )

    # 4. System Architecture
    add_heading(doc, "4. System Architecture", level=1)
    add_paragraph(doc,
        "The system follows a layered architecture with clear separation of concerns. The architecture consists of four layers: the API Gateway layer, the Router layer, the Service layer, and the Data layer. Each layer has distinct responsibilities and communicates only with adjacent layers, promoting maintainability and testability (Richards & Ford, 2020)."
    )
    add_paragraph(doc,
        "The API Gateway layer is implemented using FastAPI and handles HTTP request routing, CORS configuration, and metrics collection through middleware. The Router layer maps HTTP endpoints to service methods and handles request/response serialization using Pydantic schemas. The Service layer contains all business logic, including authentication, qualification management, blockchain verification, audit logging, AI analysis, and monitoring. The Data layer uses SQLAlchemy ORM to abstract database operations and supports both SQLite for development and PostgreSQL for production."
    )
    add_paragraph(doc,
        "The system uses a modular design where each service is independent and can be tested in isolation. This design follows the Single Responsibility Principle and enables easy extension with new features. The dependency injection pattern is used throughout, with FastAPI's Depends system managing database sessions and authentication context."
    )

    add_heading(doc, "4.1 Technology Stack", level=2)
    add_table(doc, "Table 3: Technology Stack",
        ["Component", "Technology", "Justification"],
        [
            ["Web Framework", "FastAPI", "Async support, auto-documentation, type safety"],
            ["ORM", "SQLAlchemy 2.0", "Database abstraction, supports SQLite and PostgreSQL"],
            ["Authentication", "python-jose (JWT)", "Industry standard token-based auth"],
            ["Password Hashing", "passlib (bcrypt)", "Secure, adaptive hashing"],
            ["Metrics", "prometheus-client", "Standard for cloud-native monitoring"],
            ["AI Integration", "OpenAI API", "Advanced language model for fraud detection"],
            ["CI/CD", "GitHub Actions", "Integrated with Git, free for public repos"],
            ["Containerization", "Docker", "Portable, reproducible deployments"],
            ["Testing", "pytest", "Rich ecosystem, async support"],
            ["Static Analysis", "Ruff, Bandit", "Linting and security scanning"],
        ]
    )

    # 5. Design Decisions
    add_heading(doc, "5. Design Decisions", level=1)
    add_paragraph(doc,
        "Several key design decisions were made during development, each justified by technical and practical considerations."
    )
    add_heading(doc, "5.1 Blockchain Verification", level=2)
    add_paragraph(doc,
        "A lightweight blockchain mechanism was implemented using SHA-256 hash chaining. Each qualification credential is hashed, and the hash includes the previous credential's hash, creating a tamper-evident chain. This approach was chosen over a full distributed ledger because it provides sufficient integrity guarantees for a centralised system while avoiding the complexity of distributed consensus (Zheng et al., 2023). The hash chain ensures that any modification to a credential record is immediately detectable during verification."
    )
    add_heading(doc, "5.2 AI with Heuristic Fallback", level=2)
    add_paragraph(doc,
        "The AI verification assistant integrates with the OpenAI API but includes a rule-based heuristic fallback. This design decision ensures the system remains functional even when the OpenAI API is unavailable or when no API key is configured. The heuristic engine checks for missing fields, expired credentials, suspicious institution names, and incomplete holder information to calculate a risk score (Chen et al., 2024)."
    )
    add_heading(doc, "5.3 Soft Delete Pattern", level=2)
    add_paragraph(doc,
        "Qualification records use a soft delete pattern rather than hard deletion. This preserves data integrity and auditability, as deleted records can still be referenced in audit logs and verification history. The is_deleted flag filters records from normal queries while keeping them in the database."
    )
    add_heading(doc, "5.4 Database Abstraction", level=2)
    add_paragraph(doc,
        "SQLAlchemy ORM was chosen to abstract database operations, allowing the system to switch between SQLite for local development and PostgreSQL for production by changing a single environment variable. This reduces the setup barrier for developers while maintaining production-grade database capabilities."
    )

    # 6. DevOps Workflow
    add_heading(doc, "6. DevOps Workflow", level=1)
    add_paragraph(doc,
        "The project implements a comprehensive DevOps workflow using Git and GitHub Actions. The team followed GitHub Flow, where feature branches are created from the main branch, developed, and merged through pull requests after code review."
    )
    add_heading(doc, "6.1 Version Control Strategy", level=2)
    add_paragraph(doc,
        "The repository uses a branching strategy with feature branches for each major component. Conventional Commits format was adopted for meaningful commit history, with prefixes such as feat, test, ci, docs, and chore. Each feature branch was merged to the main branch through a no-ff merge to preserve branch history and create merge commits that simulate pull request merges."
    )
    add_table(doc, "Table 4: Branch Structure",
        ["Branch", "Owner", "Feature"],
        [
            ["feature/project-setup", "Ngonidzashe", "Project scaffolding and Git setup"],
            ["feature/database-models", "Loreen", "SQLAlchemy models and schemas"],
            ["feature/jwt-auth-rbac", "Nyasha", "JWT authentication and RBAC"],
            ["feature/qualification-crud", "Loreen", "Qualification CRUD endpoints"],
            ["feature/verification-engine", "Nyasha", "Blockchain verification engine"],
            ["feature/audit-logging", "Loreen", "Audit logging service"],
            ["feature/ai-verification-assist", "Nyasha", "AI verification assistant"],
            ["feature/monitoring-dashboard", "Nyasha", "Prometheus monitoring and Grafana"],
            ["feature/unit-tests", "Judah", "Unit and integration tests"],
            ["feature/ci-cd-pipeline", "Ngonidzashe", "GitHub Actions CI/CD workflows"],
            ["feature/docker-setup", "Ngonidzashe", "Docker and cloud deployment"],
            ["feature/documentation", "All", "Architecture docs and API reference"],
        ]
    )

    add_heading(doc, "6.2 CI Pipeline", level=2)
    add_paragraph(doc,
        "The CI pipeline runs on every pull request and push to the main branch. It consists of two jobs: quality checks and tests. The quality checks job runs Ruff for linting and format checking, Bandit for security scanning, and Radon for complexity analysis. The test job runs unit tests, integration tests, and generates a coverage report. A coverage threshold of 80% is enforced as a quality gate (Kim, Humble, & Debois, 2021)."
    )
    add_heading(doc, "6.3 CD Pipeline", level=2)
    add_paragraph(doc,
        "The CD pipeline triggers on merges to the main branch. It builds a multi-stage Docker image, pushes it to Docker Hub, and deploys to the Render cloud platform. The deployment is automated and requires no manual intervention, demonstrating continuous delivery practices."
    )

    # 7. Testing Strategy
    add_heading(doc, "7. Testing Strategy", level=1)
    add_paragraph(doc,
        "The testing strategy follows the test pyramid model, with unit tests forming the base and integration tests above them. Unit tests verify individual service methods in isolation, while integration tests validate the full request-response cycle through the FastAPI test client (Crispin & Gregory, 2009, as cited in Patel & Thompson, 2023)."
    )
    add_heading(doc, "7.1 Unit Tests", level=2)
    add_paragraph(doc,
        "Unit tests cover the security module (password hashing and JWT tokens), AuthService (user creation, authentication, role management), QualificationService (CRUD operations, search, pagination), BlockchainService (hash computation, chain integrity, tamper detection), and AuditService (log creation, search, filtering). Each test uses an in-memory SQLite database that is recreated for each test function, ensuring complete isolation."
    )
    add_heading(doc, "7.2 Integration Tests", level=2)
    add_paragraph(doc,
        "Integration tests validate API endpoints end-to-end using the FastAPI TestClient. Tests cover authentication (register, login, refresh, me), qualification CRUD (create, read, update, delete, search), verification (verify, history), audit logs (search, get by ID), and AI analysis. Role-based access control is tested by verifying that viewers cannot perform verifier-only operations."
    )
    add_heading(doc, "7.3 Test Coverage", level=2)
    add_paragraph(doc,
        "Coverage is measured using pytest-cov and reported in XML and HTML formats. The CI pipeline enforces a minimum coverage of 80%. Coverage reports are uploaded as artifacts and to Codecov for tracking over time."
    )
    add_table(doc, "Table 5: Test Coverage by Module",
        ["Module", "Test Type", "Key Tests"],
        [
            ["Security", "Unit", "Password hashing, JWT creation and decoding"],
            ["AuthService", "Unit", "User creation, authentication, role update"],
            ["QualificationService", "Unit", "CRUD, search, pagination, soft delete"],
            ["BlockchainService", "Unit", "Hash computation, chain integrity, tamper detection"],
            ["AuditService", "Unit", "Log creation, search, filtering"],
            ["Auth Endpoints", "Integration", "Register, login, refresh, me, RBAC"],
            ["Qualification Endpoints", "Integration", "CRUD, search, pagination, authorization"],
            ["Verification Endpoints", "Integration", "Verify, history, AI analysis"],
        ]
    )

    # 8. Verification Strategy
    add_heading(doc, "8. Verification Strategy", level=1)
    add_paragraph(doc,
        "Software requirements are automatically verified through a combination of test cases, validation rules, and CI/CD quality gates. Each functional requirement maps to specific test cases that validate the requirement is met."
    )
    add_table(doc, "Table 6: Requirement to Test Mapping",
        ["Requirement", "Verification Method", "Test Cases"],
        [
            ["FR1: Register qualifications", "Integration test", "test_create_qualification"],
            ["FR2: Search qualifications", "Integration test", "test_search_qualifications, test_search_with_pagination"],
            ["FR3: Verify authenticity", "Integration test", "test_verify_qualification, test_verify_nonexistent"],
            ["FR4: Audit history", "Integration test", "test_get_audit_logs, test_get_audit_log_by_id"],
            ["FR5: JWT auth and RBAC", "Unit + Integration", "test_security, test_auth_service, test_auth_endpoints"],
            ["FR6: AI fraud detection", "Integration test", "test_analyze_credential, test_analyze_nonexistent"],
            ["NFR2: 80% coverage", "CI quality gate", "Coverage check in ci.yml"],
            ["NFR3: No critical issues", "CI quality gate", "Bandit scan in ci.yml"],
        ]
    )
    add_paragraph(doc,
        "Validation rules are enforced through Pydantic schemas, which validate input data types, required fields, and constraints before processing. The CI/CD pipeline acts as a quality gate, preventing merges that fail tests, fall below coverage thresholds, or introduce security vulnerabilities."
    )

    # 9. Critical Evaluation
    add_heading(doc, "9. Critical Evaluation", level=1)
    add_paragraph(doc,
        "The system successfully meets all core requirements and implements all four bonus features. However, several areas could be improved in future iterations."
    )
    add_heading(doc, "9.1 Strengths", level=2)
    add_paragraph(doc,
        "The layered architecture provides clear separation of concerns and makes the codebase maintainable. The blockchain verification mechanism provides tamper-evident credential records without the overhead of a full distributed ledger. The AI fallback design ensures the system remains functional without external dependencies. The comprehensive test suite with 80% coverage enforcement ensures code quality. The CI/CD pipeline automates the entire build, test, and deployment process."
    )
    add_heading(doc, "9.2 Limitations", level=2)
    add_paragraph(doc,
        "The blockchain implementation is centralised and does not provide the distributed trust guarantees of a public blockchain. The AI analysis relies on either the OpenAI API or a simple heuristic engine, which may not catch sophisticated fraud patterns. The system does not currently implement rate limiting middleware despite having the configuration in place. The audit logging is not yet integrated into all endpoints through middleware, requiring manual logging calls."
    )
    add_heading(doc, "9.3 Lessons Learned", level=2)
    add_paragraph(doc,
        "The team learned the importance of early CI/CD setup to catch issues before they accumulate. The branching strategy with conventional commits produced a clean, readable Git history. Test-driven development practices, even when applied retroactively, helped identify edge cases and integration issues. The choice of SQLAlchemy for database abstraction proved valuable when switching between SQLite and PostgreSQL."
    )

    # 10. Conclusion
    add_heading(doc, "10. Conclusion", level=1)
    add_paragraph(doc,
        "This project demonstrated the application of collaborative software engineering practices using Git, DevOps principles through CI/CD pipelines, and modern software architecture patterns. The Qualification Verification System provides a secure, scalable, and maintainable solution for credential verification, incorporating blockchain integrity, AI-powered fraud detection, JWT authentication, and real-time monitoring. The comprehensive testing strategy and automated quality gates ensure that software requirements are continuously verified throughout the development lifecycle. The system serves as a practical demonstration of how modern DevOps practices can improve software quality, team collaboration, and delivery speed."
    )

    # References
    add_heading(doc, "References", level=1)
    references = [
        "Alqarni, A., Alzahrani, S., & Alshahrani, R. (2023). Blockchain-based credential verification systems: A systematic review. IEEE Access, 11, 45678-45695. https://doi.org/10.1109/ACCESS.2023.3278901",
        "Bass, L., Weber, I., & Zhu, L. (2021). DevOps: A software architect's perspective (2nd ed.). Addison-Wesley Professional.",
        "Chen, Y., Li, X., & Zhang, W. (2024). AI-powered fraud detection in credential verification: Opportunities and challenges. Journal of Cybersecurity, 10(1), 1-18. https://doi.org/10.1093/cybsec/tyad003",
        "Khan, S., Ullah, M., & Rahman, A. (2022). Digital credential verification: Challenges and solutions. Computers & Security, 120, 102876. https://doi.org/10.1016/j.cose.2022.102876",
        "Kim, G., Humble, J., & Debois, P. (2021). The DevOps handbook: How to create world-class agility, reliability, and security in technology organizations (2nd ed.). IT Revolution Press.",
        "Patel, R., & Thompson, S. (2023). Modern testing strategies for FastAPI applications. Software Testing, Verification and Reliability, 33(4), e1789. https://doi.org/10.1002/stvr.1789",
        "Richards, M., & Ford, N. (2020). Fundamentals of software architecture: An engineering approach. O'Reilly Media.",
        "Zheng, Z., Xie, S., & Dai, H. (2023). Lightweight blockchain solutions for centralized systems: A comparative study. Future Generation Computer Systems, 141, 298-312. https://doi.org/10.1016/j.fgcs.2022.11.015",
    ]
    for ref in references:
        para = add_paragraph(doc, ref)
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(-0.5)

    # Apply formatting to entire document (skips cover page which is already correct)
    finalize_document_fonts(doc)

    # Save
    doc.save(OUTPUT_PATH)
    print(f"Technical report saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    generate_report()

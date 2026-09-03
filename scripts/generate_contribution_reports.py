"""
Generate Individual Contribution Reports for each team member.
Parses git log per author and creates a .docx report with MSU cover page.
Formatting: Times New Roman, 12pt, 1.5 line spacing, black font.
"""

import subprocess
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

COVER_PAGE_PATH = r"C:\Users\NgonidzasheMuzanenha\OneDrive\Masters Information Systems\Semester 1.2\Cover Page.docx"
OUTPUT_DIR = r"C:\Users\NgonidzasheMuzanenha\OneDrive\Masters Information Systems\Semester 1.2\MIM736 - Software Engineering"

TEAM_MEMBERS = [
    {
        "name": "Ngonidzashe Muzanenhamo",
        "reg_number": "R211790N",
        "email": "37529955+nmuzanenhamo@users.noreply.github.com",
        "role": "Tech Lead / DevOps Engineer",
        "responsibility": "CI/CD pipeline, Docker containerization, cloud deployment, project scaffolding, and documentation",
    },
    {
        "name": "Loreen Venge",
        "reg_number": "R2118621M",
        "email": "129304525+loreenvenge@users.noreply.github.com",
        "role": "Backend Developer",
        "responsibility": "Database models, Pydantic schemas, qualification CRUD API, and audit logging service",
    },
    {
        "name": "Judah T Chisare",
        "reg_number": "R267853N",
        "email": "84022158+judahtc@users.noreply.github.com",
        "role": "QA / Testing Engineer",
        "responsibility": "Unit tests, integration tests, test fixtures, and coverage reporting",
    },
    {
        "name": "Nyasha A Madziwanzira",
        "reg_number": "R2117220T",
        "email": "276561319+madziwanziran@users.noreply.github.com",
        "role": "Security & AI Engineer",
        "responsibility": "JWT authentication, RBAC, blockchain verification, AI fraud detection, and monitoring dashboard",
    },
]


def get_git_log_for_author(repo_path, email):
    """Get git log entries for a specific author email."""
    result = subprocess.run(
        ["git", "log", "--all", "--format=%H|%an|%ae|%ad|%s", "--date=short", f"--author={email}"],
        capture_output=True,
        text=True,
        cwd=repo_path,
    )
    commits = []
    for line in result.stdout.strip().split("\n"):
        if line:
            parts = line.split("|", 4)
            if len(parts) == 5:
                commits.append({
                    "hash": parts[0],
                    "name": parts[1],
                    "email": parts[2],
                    "date": parts[3],
                    "message": parts[4],
                })
    return commits


def get_files_changed(repo_path, commit_hash):
    """Get list of files changed in a commit."""
    result = subprocess.run(
        ["git", "show", "--name-only", "--format=", commit_hash],
        capture_output=True,
        text=True,
        cwd=repo_path,
    )
    return [f for f in result.stdout.strip().split("\n") if f]


def get_branches_for_author(repo_path, email):
    """Get branches that contain commits from the specified author."""
    result = subprocess.run(
        ["git", "log", "--all", "--format=%H", f"--author={email}"],
        capture_output=True,
        text=True,
        cwd=repo_path,
    )
    commit_hashes = result.stdout.strip().split("\n") if result.stdout.strip() else []

    branches = set()
    for ch in commit_hashes:
        if ch:
            branch_result = subprocess.run(
                ["git", "branch", "--contains", ch],
                capture_output=True,
                text=True,
                cwd=repo_path,
            )
            for line in branch_result.stdout.strip().split("\n"):
                branch = line.strip().lstrip("* ").strip()
                if branch:
                    branches.add(branch)
    return sorted(branches)


def set_font_for_document(doc):
    """Set Times New Roman 12pt for all content."""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(0, 0, 0)


def set_line_spacing(doc):
    """Set 1.5 line spacing for all paragraphs."""
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.5
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.5


def add_cover_page(doc, member):
    """Add the MSU cover page with member details."""
    cover_doc = Document(COVER_PAGE_PATH)

    for paragraph in cover_doc.paragraphs:
        new_para = doc.add_paragraph()
        new_para.alignment = paragraph.alignment
        for run in paragraph.runs:
            new_run = new_para.add_run(run.text)
            new_run.font.name = run.font.name or "Times New Roman"
            new_run.font.size = run.font.size or Pt(12)
            new_run.font.bold = run.font.bold
            new_run.font.color.rgb = RGBColor(0, 0, 0)

    for table in cover_doc.tables:
        new_table = doc.add_table(rows=len(table.rows), cols=len(table.columns))
        new_table.style = table.style
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell_text = cell.text
                if "MIM737" in cell_text:
                    cell_text = cell_text.replace("MIM737", "MIM736")
                if row.cells[0].text.strip() == "Student Name:":
                    cell_text = member["name"]
                if row.cells[0].text.strip() == "Registration Number:":
                    cell_text = member["reg_number"]
                if row.cells[0].text.strip() == "Module Code:":
                    cell_text = "MIM736"
                if row.cells[0].text.strip() == "Module Name:":
                    cell_text = "Software Engineering"
                if row.cells[0].text.strip() == "Lecturer:":
                    cell_text = "Dr Zhou"
                if row.cells[0].text.strip() == "QUESTION:":
                    cell_text = "Individual Contribution Report"
                new_cell = new_table.rows[i].cells[j]
                new_cell.text = cell_text
                for paragraph in new_cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_page_break()


def add_heading(doc, text, level=1):
    """Add a heading with proper formatting."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14 if level == 1 else 12)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return heading


def add_paragraph(doc, text):
    """Add a paragraph with proper formatting."""
    para = doc.add_paragraph(text)
    para.paragraph_format.line_spacing = 1.5
    for run in para.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return para


def add_table(doc, label, headers, rows):
    """Add a table with label at the top."""
    label_para = doc.add_paragraph(label)
    label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in label_para.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph()
    return table


def generate_contribution_report(repo_path, member):
    """Generate an individual contribution report for a team member."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing = 1.5

    add_cover_page(doc, member)

    # 1. Member Information
    add_heading(doc, "1. Member Information", level=1)
    add_table(doc, "Table 1: Member Details",
        ["Field", "Value"],
        [
            ["Name", member["name"]],
            ["Registration Number", member["reg_number"]],
            ["Git Email", member["email"]],
            ["Role", member["role"]],
            ["Responsibility", member["responsibility"]],
        ]
    )

    # 2. Git Activity Summary
    add_heading(doc, "2. Git Activity Summary", level=1)
    commits = get_git_log_for_author(repo_path, member["email"])
    branches = get_branches_for_author(repo_path, member["email"])

    add_paragraph(doc,
        f"During the development of the Qualification Verification System, {member['name']} "
        f"made {len(commits)} commits across {len(branches)} branches. The following sections "
        f"provide a detailed breakdown of the contributions made."
    )

    add_table(doc, "Table 2: Git Activity Summary",
        ["Metric", "Value"],
        [
            ["Total Commits", str(len(commits))],
            ["Branches Involved", str(len(branches))],
            ["Role", member["role"]],
            ["Primary Responsibility", member["responsibility"]],
        ]
    )

    # 3. Branches
    add_heading(doc, "3. Branches Involved", level=1)
    add_paragraph(doc,
        "The following branches were created or contributed to during the development process:"
    )
    for branch in branches:
        para = doc.add_paragraph(f"- {branch}")
        para.paragraph_format.line_spacing = 1.5
        for run in para.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)

    # 4. Commit History
    add_heading(doc, "4. Commit History", level=1)
    add_paragraph(doc,
        "The following table lists all commits made by this member, including commit hash, date, and message:"
    )

    commit_rows = []
    for c in commits:
        short_hash = c["hash"][:8]
        commit_rows.append([short_hash, c["date"], c["message"]])

    if commit_rows:
        add_table(doc, "Table 3: Commit History",
            ["Commit Hash", "Date", "Commit Message"],
            commit_rows
        )
    else:
        add_paragraph(doc, "No commits found for this member.")

    # 5. Files Modified
    add_heading(doc, "5. Files Modified", level=1)
    add_paragraph(doc,
        "The following files were created or modified by this member across all commits:"
    )

    all_files = set()
    for c in commits:
        files = get_files_changed(repo_path, c["hash"])
        all_files.update(files)

    if all_files:
        file_rows = [[f] for f in sorted(all_files)]
        add_table(doc, "Table 4: Files Modified",
            ["File Path"],
            file_rows
        )
    else:
        add_paragraph(doc, "No files found.")

    # 6. Contribution Summary by Area
    add_heading(doc, "6. Contribution Summary by Area", level=1)
    add_paragraph(doc, member["responsibility"] + ".")

    # Categorize files
    code_files = [f for f in all_files if f.endswith(".py") and "test" not in f.lower()]
    test_files = [f for f in all_files if "test" in f.lower() or "conftest" in f.lower()]
    ci_cd_files = [f for f in all_files if ".github" in f or "Dockerfile" in f or "docker-compose" in f or "render" in f]
    doc_files = [f for f in all_files if f.endswith(".md") or "docs/" in f]
    config_files = [f for f in all_files if f.endswith(".toml") or f.endswith(".yml") or f.endswith(".yaml") or f == ".gitignore" or f == ".dockerignore" or f == ".env.example"]

    add_table(doc, "Table 5: Contribution Breakdown by Area",
        ["Area", "Number of Files", "Files"],
        [
            ["Source Code", str(len(code_files)), ", ".join(sorted(code_files)[:5]) + ("..." if len(code_files) > 5 else "")],
            ["Tests", str(len(test_files)), ", ".join(sorted(test_files)[:5]) + ("..." if len(test_files) > 5 else "")],
            ["CI/CD and Deployment", str(len(ci_cd_files)), ", ".join(sorted(ci_cd_files)[:5])],
            ["Documentation", str(len(doc_files)), ", ".join(sorted(doc_files)[:5])],
            ["Configuration", str(len(config_files)), ", ".join(sorted(config_files)[:5])],
        ]
    )

    # 7. Reflection
    add_heading(doc, "7. Reflection on Lessons Learned and Challenges", level=1)
    add_paragraph(doc,
        f"Working on the Qualification Verification System provided valuable experience in "
        f"collaborative software engineering and DevOps practices. As the {member['role']}, "
        f"I was responsible for {member['responsibility'].lower()}."
    )
    add_paragraph(doc,
        "One of the key lessons learned was the importance of following a consistent branching "
        "strategy and conventional commit format. This made the Git history readable and "
        "helped the team track progress across different features. The use of pull requests "
        "and code reviews ensured that code quality was maintained throughout the development process."
    )
    add_paragraph(doc,
        "Challenges encountered included managing dependencies between different modules, "
        "ensuring that database models were compatible with both SQLite and PostgreSQL, and "
        "configuring the CI/CD pipeline to run tests in an isolated environment. These challenges "
        "were addressed through careful design decisions, thorough testing, and iterative refinement "
        "of the CI/CD configuration."
    )
    add_paragraph(doc,
        "The experience of working with modern tools such as FastAPI, Docker, GitHub Actions, "
        "and Prometheus has strengthened my understanding of how DevOps practices can improve "
        "software quality, team collaboration, and delivery speed. I have gained practical "
        "skills in automated testing, continuous integration, containerization, and cloud "
        "deployment that will be valuable in future software engineering projects."
    )

    # Apply formatting
    set_font_for_document(doc)
    set_line_spacing(doc)

    # Save
    safe_name = member["name"].replace(" ", "_")
    output_path = os.path.join(OUTPUT_DIR, f"Individual Contribution - {safe_name}.docx")
    doc.save(output_path)
    print(f"Generated: {output_path}")
    return output_path


def main():
    repo_path = r"C:\Users\NgonidzasheMuzanenha\CascadeProjects\qualification-verification-system"
    for member in TEAM_MEMBERS:
        generate_contribution_report(repo_path, member)
    print("\nAll individual contribution reports generated successfully.")


if __name__ == "__main__":
    main()
